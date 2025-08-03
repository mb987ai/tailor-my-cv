from flask import Flask, render_template, request, send_file, session
from groq import Groq
import os
import io
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = "secret123"  # Needed for session
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))


def extract_text_from_docx(file):
    try:
        document = Document(file)
        return "\n".join([para.text for para in document.paragraphs])
    except Exception:
        return ""


def save_to_docx(text):
    doc = Document()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    tone = request.form.get("tone", "human-like")
    resume_text = request.form.get("resume_text", "").strip()
    resume_file = request.files.get("resume_file")

    if not resume_text and resume_file:
        filename = secure_filename(resume_file.filename)
        if filename.endswith(".docx"):
            resume_text = extract_text_from_docx(resume_file)
        else:
            try:
                resume_text = resume_file.read().decode("utf-8")
            except Exception:
                return "❌ Could not read resume file. Please upload a .docx or valid .txt file."

    job_text = request.form.get("job_text", "").strip()
    job_file = request.files.get("job_file")

    if not job_text and job_file:
        filename = secure_filename(job_file.filename)
        if filename.endswith(".docx"):
            job_text = extract_text_from_docx(job_file)
        else:
            try:
                job_text = job_file.read().decode("utf-8")
            except Exception:
                return "❌ Could not read job description file. Please upload a .docx or valid .txt file."

    cover_text = request.form.get("cover_text", "").strip()

    if not resume_text or not job_text:
        return "❌ Please provide both resume and job description!"

    prompt = f"""
You're an expert career coach and resume writer. Please rewrite the following resume in a {tone} tone to best match the job description. Also include a tailored cover letter. If a cover letter intro is provided, include it.

Resume:
{resume_text}

Job Description:
{job_text}

Cover Letter Intro (if any):
{cover_text}
"""

    try:
        response = groq_client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[{
                "role":
                "system",
                "content":
                "You are a professional resume and cover letter writer."
            }, {
                "role": "user",
                "content": prompt
            }])

        result_text = response.choices[0].message.content.strip()

        # Save to session for download
        session['result_text'] = result_text

        return render_template("result.html", result=result_text)

    except Exception as e:
        return f"❌ An error occurred: {e}"


@app.route("/download")
def download():
    result_text = session.get('result_text', '')
    if not result_text:
        return "No result available to download."

    docx_file = save_to_docx(result_text)
    return send_file(docx_file,
                     as_attachment=True,
                     download_name="Tailored_CV_and_Cover_Letter.docx")


if __name__ == "__main__":
    app.run(debug=True)
